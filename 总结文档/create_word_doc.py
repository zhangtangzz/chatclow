from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

# 标题
title = doc.add_heading('ChatClow 项目学习记录', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('学习日期：2026-05-21')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# 一、今日概览
doc.add_heading('一、今日学习概览', level=1)
p = doc.add_paragraph()
p.add_run('今天完成了两大板块的学习：').bold = True
overview = doc.add_paragraph(style='List Bullet')
overview.add_run('安全&稳定性模块（三件套）').bold = True
overview.add_run(' —— JWT登录认证、全局异常处理、参数校验')
overview2 = doc.add_paragraph(style='List Bullet')
overview2.add_run('核心业务 Phase 2（部分）').bold = True
overview2.add_run(' —— AI模型配置模块完整全链路开发')

prog = doc.add_paragraph()
prog.add_run('当前进度：').bold = True
prog_text = doc.add_paragraph()
prog_text.add_run('基础架构+安全模块完成，核心业务进行中（约60%）')
prog_text.runs[0].font.size = Pt(10)

# 二、安全三件套
doc.add_heading('二、安全&稳定性 三件套', level=1)

# 2.1 JWT
doc.add_heading('2.1 JWT 登录认证', level=2)
p = doc.add_paragraph()
p.add_run('核心思路：用户登录成功后签发Token，后续请求通过Header携带Token进行身份验证。').italic = True

files_jwt = [
    ('util/JwtUtil.java', 'JWT工具类，负责生成/解析/验证Token'),
    ('controller/LoginController.java', '登录接口 POST /api/auth/login'),
    ('interceptor/JwtInterceptor.java', '拦截器，验证每个请求的Token'),
    ('config/WebMvcConfig.java', '注册拦截器 + 跨域配置'),
]
p2 = doc.add_paragraph()
p2.add_run('涉及文件：').bold = True
for f, desc in files_jwt:
    b = doc.add_paragraph(style='List Bullet')
    b.add_run(f).bold = True
    b.add_run(' —— ' + desc)

# 2.2 全局异常
doc.add_heading('2.2 全局异常处理', level=2)
p = doc.add_paragraph()
p.add_run('使用 @RestControllerAdvice + @ExceptionHandler 统一捕获异常，返回友好提示而非堆栈信息。').italic = True

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
headers = ['异常类型', '触发场景', '返回格式']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for para in table.rows[0].cells[i].paragraphs:
        for r in para.runs:
            r.bold = True
            r.font.size = Pt(10)
exceptions_data = [
    ('MethodArgumentNotValidException', '@Valid 校验失败时', '{"code":400,"msg":"具体错误"}'),
    ('BindException', '参数绑定错误', '{"code":400,"msg":"参数绑定失败"}'),
    ('NullPointerException', '空指针异常', '{"code":500,"msg":"系统繁忙"}'),
    ('IllegalArgumentException', '非法参数异常', '{"code":400,"msg":"参数不合法"}'),
    ('Exception (兜底)', '其他所有未捕获异常', '{"code":500,"msg":"服务器内部错误"}'),
]
for ri, (exc, scene, resp) in enumerate(exceptions_data, start=1):
    cells = table.rows[ri].cells
    cells[0].text = exc
    cells[1].text = scene
    cells[2].text = resp
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(9)

# 2.3 参数校验
doc.add_heading('2.3 参数校验', level=2)
p = doc.add_paragraph()
p.add_run('基于 JSR-303 规范（Hibernate Validator 实现），通过注解约束实体字段。').italic = True

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Table Grid'
h2 = ['注解', '作用', 'User.java 使用场景']
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(10)
annots = [
    ('@NotBlank', '不能为 null/空串/纯空格', 'username, email, password'),
    ('@Email', '必须是合法邮箱格式', 'email 字段'),
    ('@Size(min, max)', '字符串长度限制', 'username(3~20), password(6~32)'),
    ('@NotNull', '不能为 null', 'role 角色'),
]
for ri, (a, u, s) in enumerate(annots, start=1):
    c = table2.rows[ri].cells
    c[0].text = a; c[1].text = u; c[2].text = s
    for cell in c:
        for p in cell.paragraphs:
            for r in p.runs: r.font.size = Pt(9)

pn = doc.add_paragraph()
pn.add_run('关键点：').bold = True
pn.add_run('Controller 接口必须用 @RequestBody @Valid 才能触发校验！')

# 三、AI模型配置
doc.add_heading('三、AI 模型配置模块（核心业务）', level=1)
p = doc.add_paragraph()
p.add_run('这是从"增删改查练习"迈向"真正AI平台"的第一步。管理可用的AI模型。').italic = True

doc.add_heading('数据库设计 - chatclow_ai_model 表', level=3)
table3 = doc.add_table(rows=8, cols=4)
table3.style = 'Table Grid'
dh = ['字段名', '类型', '说明', '举例']
for i, h in enumerate(dh):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(10)
db_rows = [
    ('id', 'BIGINT AUTO_INCREMENT', '主键ID', '1'),
    ('name', 'VARCHAR(50)', '模型显示名称', 'GPT-4o'),
    ('provider', 'VARCHAR(30)', '提供商', 'openai / qwen / deepseek'),
    ('model_code', 'VARCHAR(50)', '调用时的编码', 'gpt-4o / qwen-plus'),
    ('api_url', 'VARCHAR(255)', 'API接口地址', 'https://api.openai.com/v1/chat/completions'),
    ('api_key', 'VARCHAR(255)', 'API密钥', 'sk-xxxxx...'),
    ('status', 'TINYINT DEFAULT 1', '状态：0禁用 1启用', '1'),
]
for ri, row in enumerate(db_rows, start=1):
    for di, val in enumerate(row):
        table3.rows[ri].cells[di].text = val
        for p in table3.rows[ri].cells[di].paragraphs:
            for r in p.runs: r.font.size = Pt(9)

doc.add_heading('全链路开发文件', level=3)
chain_files = [
    ('Entity', 'AiModel.java', '@TableName + @NotBlank 校验注解'),
    ('Mapper', 'AiModelMapper.java', '@Mapper + BaseMapper<AiModel>'),
    ('Service接口', 'AiModelService.java', 'add/listEnabled/getById/update/delete/toggleStatus'),
    ('Service实现', 'AiModelServiceImpl.java', 'LambdaQueryWrapper + toggleStatus切换逻辑'),
    ('Controller', 'AiModelController.java', '6个RESTful接口，带@Valid校验'),
]
t4 = doc.add_table(rows=len(chain_files)+1, cols=3)
t4.style = 'Table Grid'
for i, h in enumerate(['层级', '文件', '要点']):
    t4.rows[0].cells[i].text = h
    for p in t4.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(10)
for ri, (layer, f, pt) in enumerate(chain_files, start=1):
    t4.rows[ri].cells[0].text = layer
    t4.rows[ri].cells[1].text = f
    t4.rows[ri].cells[2].text = pt
    for c in t4.rows[ri].cells:
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9)

doc.add_heading('API 接口清单', level=3)
apis = [
    ('POST', '/api/model/add', '新增模型(@Valid校验)'),
    ('GET', '/api/model/list', '查询所有启用的模型'),
    ('GET', '/api/model/{id}', '根据ID查询'),
    ('PUT', '/api/model/update', '更新模型(@Valid校验)'),
    ('DELETE', '/api/model/delete/{id}', '删除模型'),
    ('PUT', '/api/model/status/{id}', '切换启用/禁用状态'),
]
t5 = doc.add_table(rows=len(apis)+1, cols=3)
t5.style = 'Table Grid'
for i, h in enumerate(['方法', '路径', '功能']):
    t5.rows[0].cells[i].text = h
    for p in t5.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True; r.font.size = Pt(10)
for ri, (method, path, func) in enumerate(apis, start=1):
    t5.rows[ri].cells[0].text = method
    t5.rows[ri].cells[1].text = path
    t5.rows[ri].cells[2].text = func
    for c in t5.rows[ri].cells:
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9)

# 四、踩坑记录
doc.add_heading('四、踩坑记录 & 解决方案', level=1)
bugs = [
    ('Spring Boot 2.6 拦截器白名单失效', '登录接口返回401未登录', 'Spring Boot 2.6默认路径匹配策略改变', 'application.yml添加 ant_path_matcher'),
    ('Java 17 JAXB 缺失', '启动报错 ClassNotFoundException', 'jjwt 0.9.1使用javax.xml.bind，Java 9+移除了', 'pom.xml加 jaxb-api + jaxb-runtime依赖'),
    ('Bearer 大小写问题', '带Token仍返回401', 'Header中写了小写bearer，必须大写Bearer', 'Authorization: Bearer <token>(B大写)'),
    ('Token放错了位置', 'ApiPost带了Token还是401', 'Token写在Body标签页，应在Header标签页', 'ApiPost -> Headers标签页添加Authorization'),
    ('R.ok()单参数类型冲突', 'R.ok("成功!") 返回R<Void>报错', '单参数版本把String当data处理，String!=Void', '改用 R.ok("成功!", null) 双参数版本'),
    ('Mapper文件名写错', '新建Mapper后IDE报错', '文件名写成AiModel.java(和实体类重名)', '重命名为 AiModelMapper.java'),
]

for idx, (title, problem, reason, fix) in enumerate(bugs, start=1):
    bp = doc.add_paragraph()
    bp.add_run('坑 %d：%s' % (idx, title)).bold = True
    for label, text in [('现象：', problem), ('原因：', reason), ('解决：', fix)]:
        b = doc.add_paragraph(style='List Bullet')
        b.add_run(label).bold = True
        b.add_run(text)
    doc.add_paragraph()

# 五、知识点总结
doc.add_heading('五、今日核心知识点', level=1)
knowledge = [
    ('JWT (JSON Web Token)', '无状态认证方式，包含Header.Payload.Signature三部分。服务端用密钥签名验证，无需存储会话。'),
    ('HandlerInterceptor', 'Spring MVC拦截器接口，preHandle在Controller执行前拦截请求，适合做鉴权和日志。'),
    ('WebMvcConfigurer', 'Spring MVC配置接口，用于注册拦截器、配置跨域等，配合@Configuration使用。'),
    ('@RestControllerAdvice', '全局控制器增强，配合@ExceptionHandler实现统一异常处理，RESTful API标准做法。'),
    ('JSR-303 Bean Validation', 'Java标准参数校验规范，常用@NotBlank/@NotNull/@Size/@Email/@Pattern，需搭配@Valid使用。'),
    ('LambdaQueryWrapper', 'MyBatis-Plus条件构造器，用方法引用构建查询条件，避免手写字段名字符串出错。'),
    ('@RequestBody vs @RequestParam', '@RequestBody接JSON对象(+@Valid)；@RequestParam接单个参数或表单字段。'),
    ('Spring Boot 版本差异', '2.x用javax.validation；3.x改为jakarta.validation；2.6.x需配ant_path_matcher。'),
]
for title, content in knowledge:
    kp = doc.add_paragraph()
    kp.add_run('> %s' % title).bold = True
    kc = doc.add_paragraph(content)
    kc.paragraph_format.left_indent = Cm(0.5)

# 六、项目结构
doc.add_heading('六、项目当前完整结构', level=1)
structure = '''chatclow_api/
+-- ChatclowApplication.java          [OK] 启动类
+-- common/R.java                     [OK] 统一响应封装
+-- config/WebMvcConfig.java          [OK] 拦截器注册 + 跨域
+-- exception/GlobalExceptionHandler   [OK] 全局异常处理
+-- entity/
|   +-- User.java                     [OK] 用户实体(含校验注解)
|   +-- AiModel.java                  [OK] AI模型实体(含校验注解)
|   +-- AgentConversation.java        [OK] 会话实体
|   +-- AgentConversationRecord.java  [OK] 会话记录实体
+-- mapper/                           (4个 Mapper接口)
+-- service/                          (4组 Service接口+实现)
+-- controller/                       (5个 Controller)
+-- interceptor/JwtInterceptor.java   [OK] Token拦截
+-- util/JwtUtil.java                 [OK] JWT工具'''
sp = doc.add_paragraph()
sr = sp.add_run(structure)
sr.font.name = 'Consolas'
sr.font.size = Pt(9)

# 七、下一步计划
doc.add_heading('七、下次学习计划', level=1)
plans = [
    ('Agent 智能体管理模块', '建表 chatclow_agent -> Entity/Mapper/Service/ServiceImpl/Controller 全链路。每个Agent有名称、描述、系统提示词(system prompt)、绑定的模型ID。', '中等'),
    ('AI 对话接口（核心亮点！）', '创建对话Service，调用外部LLM API，使用SSE实现流式返回，消息存入conversation_record表。这是整个项目最核心的功能！', '较高'),
]
for title, desc, diff in plans:
    pp = doc.add_paragraph()
    pp.add_run('%s [难度:%s]' % (title, diff)).bold = True
    dp = doc.add_paragraph(desc)
    dp.paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()

final = doc.add_paragraph()
final_r = final.add_run('目标：做完这两块，项目就从"增删改查练习"升级为"真正的AI对话平台"！')
final_r.bold = True
final_r.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)

output_path = r'C:\Users\张亮亮\Desktop\ChatClow学习记录-2026-05-21.docx'
doc.save(output_path)
print('DONE: ' + output_path)
